#!/usr/bin/env python3
"""
Create comprehensive sample documents for the RAG system.
Generates detailed multi-paragraph documents for testing.
"""

from fpdf import FPDF
from docx import Document
import os

def create_comprehensive_pdf(filename, title, sections):
    """Create a comprehensive PDF document."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)

    # Title
    pdf.cell(0, 10, title, 0, 1, 'C')
    pdf.ln(10)

    pdf.set_font("Arial", '', 12)

    for section_title, paragraphs in sections:
        # Section header
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 8, section_title, 0, 1)
        pdf.ln(3)

        # Section content
        pdf.set_font("Arial", '', 12)
        for paragraph in paragraphs:
            lines = pdf.multi_cell(0, 6, paragraph)
            pdf.ln(3)

        pdf.ln(5)

    pdf.output(filename)

def create_comprehensive_docx(filename, title, sections):
    """Create a comprehensive DOCX document."""
    doc = Document()
    doc.add_heading(title, 0)

    for section_title, paragraphs in sections:
        doc.add_heading(section_title, level=1)

        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

    doc.save(filename)

def get_hr_policy_content():
    """Generate comprehensive HR policy content."""
    return [
        ("Introduction", [
            "This comprehensive Human Resources Policy Manual serves as the foundation for TechCorp Solutions' commitment to creating a positive, productive, and inclusive work environment. Our policies are designed to protect both employees and the organization while promoting fairness, respect, and professional growth.",
            "These policies apply to all employees, contractors, and temporary workers engaged by TechCorp Solutions. Management has the responsibility to ensure compliance with these policies and to communicate them effectively to their teams."
        ]),
        ("Equal Employment Opportunity", [
            "TechCorp Solutions is committed to providing equal employment opportunities to all qualified individuals regardless of race, color, religion, sex, national origin, age, disability, sexual orientation, gender identity, or veteran status. We maintain a zero-tolerance policy for discrimination in all aspects of employment.",
            "Our recruitment process is designed to attract diverse talent and ensure fair evaluation of all candidates. Interview panels include diverse representation, and all hiring decisions are based on job-related qualifications and performance.",
            "We regularly conduct diversity and inclusion training for all employees and managers to foster an inclusive workplace culture where everyone feels valued and respected."
        ]),
        ("Workplace Harassment", [
            "Harassment of any kind is strictly prohibited and will not be tolerated under any circumstances. This includes sexual harassment, bullying, intimidation, and any form of unwelcome conduct that creates a hostile work environment.",
            "Employees who experience or witness harassment should report it immediately to their supervisor, HR, or through our anonymous reporting hotline. All reports are investigated promptly and confidentially.",
            "Retaliation against anyone who reports harassment is strictly prohibited and will result in disciplinary action. We are committed to maintaining a safe and respectful workplace for all employees."
        ]),
        ("Performance Management", [
            "Our performance management system is designed to help employees develop their skills, achieve their goals, and contribute to company success. Regular feedback, goal setting, and performance reviews are essential components of this system.",
            "Employees receive quarterly performance check-ins with their managers to discuss progress, challenges, and development opportunities. Annual performance reviews include compensation adjustments and career planning discussions.",
            "Performance improvement plans are developed for employees who are not meeting expectations, providing clear objectives, timelines, and support resources to help them succeed."
        ]),
        ("Compensation and Benefits", [
            "TechCorp Solutions offers competitive compensation packages that include base salary, performance bonuses, and equity participation for eligible employees. Our compensation structure is regularly benchmarked against industry standards.",
            "Benefits include comprehensive health insurance, dental and vision coverage, retirement savings plans with company matching, paid time off, and professional development allowances. We regularly review our benefits package to ensure it meets employee needs.",
            "Salary increases are determined through our annual performance review process and merit-based compensation adjustments. Market adjustments are made periodically to ensure our compensation remains competitive."
        ]),
        ("Professional Development", [
            "We believe in investing in our employees' growth and development. Professional development opportunities include training programs, conference attendance, certification support, and tuition reimbursement for job-related education.",
            "Employees are encouraged to create individual development plans with their managers, identifying skills they want to develop and career goals they wish to pursue. The company provides resources and time to support these development activities.",
            "Leadership development programs are available for high-potential employees, preparing them for future leadership roles within the organization."
        ])
    ]

def get_it_security_content():
    """Generate comprehensive IT security content."""
    return [
        ("Security Overview", [
            "Information security is fundamental to TechCorp Solutions' operations and our clients' trust. This comprehensive security framework protects sensitive data, ensures regulatory compliance, and mitigates cyber threats through layered security controls and employee awareness.",
            "Our security program follows industry best practices including NIST, ISO 27001, and SOC 2 standards. Regular audits and assessments ensure our security measures remain effective against evolving threats."
        ]),
        ("Access Management", [
            "Access to systems and data is granted based on the principle of least privilege, ensuring employees have access only to the information necessary for their job functions. Role-based access control (RBAC) is implemented across all systems.",
            "Multi-factor authentication (MFA) is required for all remote access and privileged accounts. Password policies require complex passwords with regular rotation and account lockout after failed attempts.",
            "Access reviews are conducted quarterly to ensure permissions remain appropriate as employees change roles or leave the company. Immediate deprovisioning occurs when employees depart."
        ]),
        ("Data Protection", [
            "All sensitive data is classified and protected according to its sensitivity level. Encryption is used for data at rest and in transit using industry-standard algorithms. Data loss prevention (DLP) tools prevent unauthorized data exfiltration.",
            "Backup and recovery procedures ensure business continuity and data availability. Regular backups are tested, and recovery time objectives are established for critical systems.",
            "Data retention policies define how long different types of data are kept, balancing legal requirements with storage efficiency. Secure deletion methods are used when data reaches the end of its retention period."
        ]),
        ("Network Security", [
            "Our network architecture includes multiple layers of defense including firewalls, intrusion detection systems, and network segmentation. Regular vulnerability scanning and penetration testing identify and address security weaknesses.",
            "Wireless networks are secured with WPA3 encryption and network access control. Guest networks are isolated from internal networks to prevent unauthorized access.",
            "Remote access is secured through VPN connections with endpoint verification. All remote devices must comply with security standards before connecting to corporate networks."
        ]),
        ("Endpoint Security", [
            "All company devices are protected with comprehensive endpoint security solutions including antivirus, endpoint detection and response (EDR), and device encryption. Regular security updates and patches are applied automatically.",
            "Mobile device management (MDM) ensures company-issued mobile devices comply with security policies. Personal devices used for work (BYOD) must meet minimum security requirements.",
            "Security awareness training is mandatory for all employees, covering topics like phishing recognition, password security, and safe internet practices. Simulated phishing exercises test employee readiness."
        ]),
        ("Incident Response", [
            "Our incident response plan provides structured procedures for detecting, responding to, and recovering from security incidents. The incident response team is trained and available 24/7 for critical incidents.",
            "All security incidents must be reported immediately through established channels. Initial triage determines the severity and appropriate response level. Communication protocols ensure stakeholders are informed appropriately.",
            "Post-incident reviews analyze what happened, why it occurred, and how to prevent similar incidents. Lessons learned are incorporated into updated security measures and training programs."
        ]),
        ("Compliance and Audit", [
            "Regular security audits and compliance assessments ensure our security program meets regulatory requirements and industry standards. External auditors provide independent verification of our security controls.",
            "Documentation of security policies, procedures, and controls is maintained and regularly updated. Version control and approval processes ensure documentation remains current and accurate.",
            "Third-party vendors and partners must meet our security requirements before being granted access to our systems or data. Regular security assessments are conducted for critical vendors."
        ])
    ]

def get_benefits_content():
    """Generate comprehensive benefits content."""
    return [
        ("Health and Wellness", [
            "TechCorp Solutions prioritizes employee health and wellness through comprehensive medical coverage, preventive care, and wellness programs. Our health insurance plans provide extensive coverage for medical, dental, and vision services.",
            "Preventive care services are fully covered, including annual physicals, vaccinations, and screening tests. Wellness programs include gym memberships, smoking cessation support, and mental health resources.",
            "Telemedicine services provide convenient access to healthcare professionals 24/7. Employees can consult with doctors virtually for non-emergency medical needs, reducing wait times and improving access to care."
        ]),
        ("Financial Security", [
            "Our retirement savings program includes a 401(k) plan with generous company matching contributions. Employees can contribute up to IRS limits with company matching up to 6% of salary.",
            "Financial planning resources include access to certified financial advisors, retirement planning workshops, and investment education. Employees can make informed decisions about their financial future.",
            "Life insurance and disability coverage protect employees and their families. Basic life insurance is provided at no cost, with options for supplemental coverage based on salary multiples."
        ]),
        ("Work-Life Balance", [
            "Generous paid time off includes 20 days of vacation annually, increasing with tenure, plus 12 paid holidays and sick leave. Employees are encouraged to take time for rest and personal matters.",
            "Flexible work arrangements support work-life balance, including remote work options, flexible hours, and compressed workweeks. Parental leave provides 12 weeks of paid leave for new parents.",
            "Employee assistance programs offer confidential counseling for personal and work-related challenges. Legal and financial consultation services are also available to employees and their families."
        ]),
        ("Career Development", [
            "Professional development is supported through tuition reimbursement, conference attendance, and training programs. Employees can pursue certifications and advanced degrees relevant to their roles.",
            "Leadership development programs prepare high-potential employees for future leadership roles. Mentoring programs connect employees with experienced leaders for career guidance.",
            "Career transition support includes resume writing, interview coaching, and job search assistance for employees seeking new opportunities, either within or outside the company."
        ]),
        ("Family Support", [
            "Family support benefits include backup childcare, elder care resources, and adoption assistance. Employees can access emergency childcare services when needed.",
            "School support programs help employees balance work and family responsibilities. College planning resources assist with children's education expenses and planning.",
            "Domestic partner benefits extend coverage to same-sex partners and spouses, ensuring all employees receive equal benefits regardless of marital status."
        ]),
        ("Recognition and Perks", [
            "Employee recognition programs celebrate achievements and milestones. Service awards, performance bonuses, and peer recognition programs motivate and reward outstanding contributions.",
            "Perks include commuter benefits, company-sponsored social events, and discounts on products and services. TechCorp branded merchandise and technology discounts are popular employee favorites.",
            "Volunteer time off encourages community involvement with paid time for volunteer activities. Company matching for charitable donations supports employees' philanthropic efforts."
        ])
    ]

def get_remote_work_content():
    """Generate comprehensive remote work content."""
    return [
        ("Remote Work Philosophy", [
            "TechCorp Solutions embraces remote work as a strategic advantage, enabling us to attract top talent worldwide while reducing overhead costs and environmental impact. Our remote-first approach prioritizes flexibility, productivity, and employee satisfaction.",
            "Research shows that remote work can increase productivity, reduce turnover, and improve work-life balance. We design our policies and tools to maximize these benefits while maintaining strong team collaboration and company culture."
        ]),
        ("Eligibility and Setup", [
            "Most roles at TechCorp are eligible for remote work arrangements, with some client-facing positions requiring occasional office presence. New hires can start remotely if their role supports it.",
            "Remote employees receive stipends for home office setup including ergonomic furniture, high-speed internet, and necessary equipment. We provide company-issued laptops, monitors, and peripherals to ensure productivity.",
            "Home office assessments help employees create optimal workspaces. Ergonomics training and equipment recommendations prevent strain and promote long-term health."
        ]),
        ("Communication and Collaboration", [
            "Effective communication is essential for remote teams. We use multiple channels including video conferencing, instant messaging, and project management tools to keep everyone connected.",
            "Regular virtual meetings, team check-ins, and one-on-one sessions maintain personal connections. Informal virtual coffee chats and social events build team camaraderie.",
            "Documentation and transparency are emphasized. Decisions, processes, and updates are documented to ensure all team members have access to important information."
        ]),
        ("Productivity and Performance", [
            "Productivity is measured by outcomes and deliverables rather than hours worked. Flexible schedules allow employees to work when they are most productive, whether early morning or late evening.",
            "Regular feedback and performance reviews ensure alignment with company goals. Remote employees have the same opportunities for advancement and recognition as office-based staff.",
            "Time tracking and project management tools help teams coordinate across time zones. Clear expectations and deadlines ensure smooth collaboration."
        ]),
        ("Security and Compliance", [
            "Remote work security requires special attention. All employees receive training on secure remote work practices, including VPN usage, data handling, and phishing awareness.",
            "Company devices include security software and remote wipe capabilities. Personal devices used for work must meet security standards and be regularly updated.",
            "Compliance with data protection regulations is maintained through secure tools and processes. Remote employees follow the same security protocols as office staff."
        ]),
        ("Wellness and Inclusion", [
            "Remote work can sometimes feel isolating, so we prioritize employee wellness through virtual team building, mental health support, and regular check-ins.",
            "Inclusive practices ensure all remote employees feel valued and included. We celebrate diversity and provide equal opportunities regardless of location.",
            "Travel and face-to-face meetings are scheduled periodically to strengthen relationships. Company retreats and team gatherings help maintain company culture."
        ]),
        ("Tools and Resources", [
            "Remote employees have access to comprehensive digital tools including video conferencing, cloud storage, project management software, and collaboration platforms.",
            "IT support is available 24/7 for remote employees. Help desk services, device troubleshooting, and technical assistance ensure minimal downtime.",
            "Learning and development resources are fully accessible remotely. Online training, virtual conferences, and digital libraries support continuous professional growth."
        ])
    ]

def get_code_of_conduct_content():
    """Generate comprehensive code of conduct content."""
    return [
        ("Ethical Standards", [
            "Integrity forms the foundation of TechCorp Solutions' culture. We conduct business with honesty, transparency, and ethical behavior in all interactions with customers, partners, competitors, and each other.",
            "Employees are expected to make decisions that benefit the company while considering the impact on stakeholders. Conflicts of interest must be disclosed and managed appropriately.",
            "Our commitment to ethics extends beyond legal requirements. We strive to do what is right, even when it's not required by law."
        ]),
        ("Respect and Inclusion", [
            "Respect for individual dignity is fundamental to our workplace. We value diversity and create an inclusive environment where everyone feels welcome and valued.",
            "Harassment, discrimination, and bullying are strictly prohibited. All employees have the right to work in an environment free from intimidation and hostility.",
            "We celebrate differences and leverage diverse perspectives to drive innovation and better decision-making."
        ]),
        ("Professional Conduct", [
            "Professional behavior is expected in all business and social interactions. Appropriate language, attire, and conduct maintain our professional reputation.",
            "Confidentiality is essential. Sensitive company and client information must be protected and shared only on a need-to-know basis.",
            "Competitive behavior should be ethical and fair. We compete vigorously but maintain integrity in all business dealings."
        ]),
        ("Compliance and Legal", [
            "Compliance with laws and regulations is mandatory. Employees must understand and follow applicable laws in their areas of work.",
            "Anti-corruption policies prohibit bribery, kickbacks, and other corrupt practices. Gifts and entertainment must be modest and appropriate.",
            "Reporting concerns is encouraged. Anonymous reporting channels are available for concerns about unethical behavior or policy violations."
        ]),
        ("Data and Information", [
            "Responsible handling of data is critical. Personal information, intellectual property, and sensitive data must be protected and used appropriately.",
            "Information security includes protecting against unauthorized access, viruses, and data breaches. Employees receive regular security training.",
            "Social media use should be professional and not damage company reputation. Personal opinions should be clearly identified as such."
        ]),
        ("Environmental Responsibility", [
            "We are committed to environmental sustainability. Employees are encouraged to reduce waste, conserve energy, and support green initiatives.",
            "Sustainable practices include recycling, reducing paper use, and choosing environmentally friendly options when available.",
            "Our environmental commitment extends to our supply chain and business partners."
        ]),
        ("Community Engagement", [
            "Community involvement strengthens our reputation and supports social good. Employees are encouraged to participate in volunteer activities.",
            "Company-sponsored community programs provide opportunities for employees to contribute to local communities and causes.",
            "Philanthropic activities align with our values and support education, technology access, and community development."
        ])
    ]

def get_holiday_policy_content():
    """Generate comprehensive holiday policy content."""
    return [
        ("Paid Time Off Overview", [
            "TechCorp Solutions provides generous paid time off to support work-life balance and employee wellness. Our PTO policy includes vacation days, holidays, and personal time that increases with tenure.",
            "PTO accrues monthly based on years of service and full-time/part-time status. Employees can view their balance and request time off through our HR portal.",
            "Advance planning is encouraged. Time off requests should be submitted at least two weeks in advance when possible, though emergency situations are accommodated."
        ]),
        ("Vacation Accrual", [
            "Full-time employees accrue 15 vacation days annually during their first three years, increasing to 20 days after three years and 25 days after seven years of service.",
            "Part-time employees accrue vacation hours proportionally based on their scheduled hours. Vacation time can be carried over up to a maximum of 30 days.",
            "Employees may cash out unused vacation time annually, though we encourage using time for rest and rejuvenation."
        ]),
        ("Holiday Schedule", [
            "TechCorp observes 10 paid holidays annually: New Year's Day, Martin Luther King Jr. Day, Memorial Day, Independence Day, Labor Day, Thanksgiving Day, Day after Thanksgiving, and Christmas Day.",
            "Floating holidays allow employees to choose additional days off based on personal or religious observances. These must be scheduled in advance and approved by management.",
            "Holiday scheduling considers business needs. Some roles may require holiday coverage with premium pay or compensatory time."
        ]),
        ("Sick Leave Policy", [
            "Employees receive 10 days of paid sick leave annually to cover illness, medical appointments, or family care needs. Sick leave accrues monthly and carries over year after year.",
            "Certification may be required for extended absences. We respect employee privacy and use certification only when necessary to protect the workplace.",
            "Employees can donate unused sick leave to colleagues facing extended medical situations through our sick leave sharing program."
        ]),
        ("Parental Leave", [
            "New parents receive 12 weeks of paid parental leave for birth, adoption, or foster care placement. This leave can be taken consecutively or intermittently.",
            "Unpaid leave extensions are available under federal and state laws. Employees can combine paid parental leave with other leave types for extended time off.",
            "Returning parents receive flexible scheduling and support to ease the transition back to work."
        ]),
        ("Bereavement Leave", [
            "Employees receive up to five days of paid bereavement leave for immediate family members and three days for extended family. Additional time may be approved for travel or special circumstances.",
            "Family is defined broadly to include spouses, children, parents, siblings, grandparents, and other close relationships.",
            "Compassionate support includes counseling services and flexible scheduling during difficult times."
        ]),
        ("Leave Administration", [
            "All leave requests are processed through our HR system with transparent approval processes. Employees receive confirmation and can track their leave balances.",
            "Intermittent leave is available for serious health conditions or family care needs. Documentation requirements vary by leave type and duration.",
            "Return-to-work procedures ensure smooth reintegration. Employees receive updates on work matters during extended absences."
        ])
    ]

def main():
    """Generate all comprehensive sample documents."""
    os.makedirs("data", exist_ok=True)

    documents = [
        ("hr_policy.pdf", "Human Resources Policy Manual", get_hr_policy_content()),
        ("it_security_guidelines.docx", "Information Technology Security Guidelines", get_it_security_content()),
        ("employee_benefits.pdf", "Employee Benefits Guide", get_benefits_content()),
        ("remote_work_policy.pdf", "Remote Work Policy and Guidelines", get_remote_work_content()),
        ("code_of_conduct.pdf", "Employee Code of Conduct", get_code_of_conduct_content()),
        ("holiday_time_off.pdf", "Holiday and Time Off Policy", get_holiday_policy_content()),
    ]

    for filename, title, content in documents:
        filepath = os.path.join("data", filename)

        if filename.endswith('.pdf'):
            create_comprehensive_pdf(filepath, title, content)
        elif filename.endswith('.docx'):
            create_comprehensive_docx(filepath, title, content)

        print(f"✅ Created: {filepath} - {len(content)} sections, {sum(len(paragraphs) for _, paragraphs in content)} total paragraphs")

if __name__ == "__main__":
    main()
