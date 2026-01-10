#!/usr/bin/env python3
"""
Generate fake enterprise documents for testing the RAG system.
Creates 6 documents with 5-12 paragraphs each in PDF and DOCX formats.
"""

from fpdf import FPDF
from docx import Document
import os

def create_pdf(filename, title, paragraphs):
    """Create a PDF document with title and paragraphs."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, title, 0, 1, 'C')
    pdf.ln(10)

    pdf.set_font("Arial", '', 12)

    for paragraph in paragraphs:
        # Handle line breaks and word wrapping
        lines = pdf.multi_cell(0, 6, paragraph)
        pdf.ln(5)

    pdf.output(filename)

def create_docx(filename, title, paragraphs):
    """Create a DOCX document with title and paragraphs."""
    doc = Document()
    doc.add_heading(title, 0)

    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    doc.save(filename)

def generate_termination_policy():
    """Generate termination policy content."""
    return [
        "This termination policy outlines the procedures and guidelines for employee separation from TechCorp Solutions. The policy ensures fair and consistent treatment of all employees during the termination process, maintaining dignity and respect throughout the transition.",
        "Termination may occur voluntarily through resignation or involuntarily through dismissal. Involuntary termination can result from performance issues, misconduct, redundancy, or organizational restructuring. All terminations must follow proper documentation and approval procedures.",
        "Employees are required to provide at least two weeks' notice for voluntary termination, except in cases of immediate resignation due to extenuating circumstances. The company reserves the right to accept shorter notice periods or request immediate departure based on role sensitivity.",
        "Upon termination, employees will receive their final paycheck within the legally mandated timeframe, typically within 7 business days. This includes all accrued but unused vacation time, pending bonuses, and any other owed compensation.",
        "Exit interviews are conducted to gather feedback about the employee's experience and identify areas for organizational improvement. Participation in exit interviews is voluntary but strongly encouraged.",
        "Company property including laptops, access cards, company phones, and confidential documents must be returned on the last day of employment. Failure to return company property may result in deductions from final paychecks.",
        "References will be provided upon request, limited to dates of employment and position held, unless otherwise authorized by the employee. The company follows standard practices for employment verification.",
        "Severance packages may be offered at the discretion of management for long-term employees or in cases of redundancy. These packages are negotiated individually and may include extended benefits or financial compensation.",
        "Employees terminated for cause may be ineligible for rehire and may forfeit certain benefits. The company maintains records of terminations to inform future hiring decisions.",
        "This policy is reviewed annually and updated as needed to comply with changing employment laws and best practices. Employees are encouraged to review the policy regularly and contact HR with any questions."
    ]

def generate_remote_work_policy():
    """Generate remote work policy content."""
    return [
        "TechCorp Solutions supports flexible work arrangements including remote work options to promote work-life balance and accommodate diverse employee needs. This policy establishes guidelines for remote work eligibility and expectations.",
        "Eligibility for remote work is determined based on job function, performance history, and managerial approval. Core functions requiring on-site presence may not be eligible for full remote arrangements.",
        "Remote employees must maintain reliable internet connectivity and a dedicated workspace that ensures privacy and minimizes distractions. The company may provide equipment stipends to support remote work setups.",
        "Core working hours are maintained regardless of location, with flexibility for start and end times within established parameters. Regular team meetings and collaboration sessions ensure continued connectivity with the organization.",
        "Productivity is measured by deliverables and results rather than hours worked. Remote employees are expected to maintain the same level of performance and accountability as on-site employees.",
        "Communication protocols include regular check-ins, use of collaboration tools, and clear expectations for response times. Video conferencing is preferred for important meetings to maintain personal connections.",
        "Security requirements include use of company-approved devices, VPN access for sensitive work, and adherence to data protection policies. Personal devices may be used only with explicit approval and proper security measures.",
        "Expense reimbursement covers reasonable home office costs including internet, utilities, and office supplies. Receipts must be submitted monthly for reimbursement processing.",
        "Performance evaluations include assessment of remote work effectiveness, including ability to collaborate virtually and maintain team cohesion. Regular feedback helps identify areas for improvement.",
        "This policy may be updated based on business needs and technological advancements. Employees are notified of changes and provided training as needed."
    ]

def generate_holiday_policy():
    """Generate holiday policy content."""
    return [
        "TechCorp Solutions provides comprehensive paid time off benefits to support employee well-being and work-life balance. This policy outlines vacation, holiday, and leave entitlements for all employees.",
        "Full-time employees accrue 15 days of paid vacation annually, increasing to 20 days after 3 years of service and 25 days after 7 years. Part-time employees receive prorated vacation based on hours worked.",
        "Vacation requests should be submitted at least two weeks in advance through the HR system. Approval is granted based on business needs and staffing requirements, with priority given to long-term planning.",
        "The company observes 10 paid holidays annually including New Year's Day, Martin Luther King Jr. Day, Memorial Day, Independence Day, Labor Day, Thanksgiving Day, and Christmas Day. Additional floating holidays may be available.",
        "Holiday scheduling considers business continuity and equal distribution across teams. Some roles may require holiday coverage with premium pay or compensatory time off.",
        "Unused vacation carries over to the next year with a maximum accumulation of 30 days. Beyond this limit, excess days may be forfeited or paid out at the company's discretion.",
        "Sick leave provides 10 days annually for illness or medical appointments. Certification may be required for extended absences. Family and medical leave follows federal and state regulations.",
        "Parental leave offers 12 weeks of paid leave for new parents, in addition to any applicable federal leave. Adoption and foster care placements are also covered under this policy.",
        "Bereavement leave provides up to 5 days for immediate family members and 3 days for extended family. Additional time may be approved for travel or special circumstances.",
        "All leave requests are confidential and handled with sensitivity. The company supports employees through life events and encourages open communication about leave needs."
    ]

def generate_code_of_conduct():
    """Generate code of conduct content."""
    return [
        "This Code of Conduct establishes the standards of behavior expected from all TechCorp Solutions employees, contractors, and representatives. It reflects our commitment to ethical business practices and professional integrity.",
        "Employees must conduct business with honesty and integrity, avoiding conflicts of interest and maintaining confidentiality of sensitive information. Personal gain should never influence professional decisions.",
        "Respect for colleagues, clients, and stakeholders is fundamental to our workplace culture. Harassment, discrimination, or bullying in any form is strictly prohibited and will result in disciplinary action.",
        "Compliance with laws and regulations is mandatory. Employees must report suspected violations through appropriate channels without fear of retaliation.",
        "Intellectual property rights are respected and protected. Company resources, including time, equipment, and information, must be used appropriately and not for personal gain.",
        "Environmental responsibility includes sustainable practices and resource conservation. Employees are encouraged to minimize waste and support green initiatives.",
        "Diversity and inclusion are core values. We celebrate differences and provide equal opportunities regardless of race, gender, age, religion, or background.",
        "Professional development is encouraged through training and education. Employees should maintain current knowledge and skills in their areas of expertise.",
        "Community involvement demonstrates our commitment to social responsibility. Volunteer activities and charitable contributions are supported when they align with company values.",
        "This Code is reviewed regularly and violations may result in disciplinary action up to and including termination. Ethical behavior is essential to our success and reputation."
    ]

def generate_security_guidelines():
    """Generate security guidelines content."""
    return [
        "Information security is critical to TechCorp Solutions' operations and reputation. These guidelines establish requirements for protecting company and client data from unauthorized access, disclosure, or modification.",
        "Access controls include unique user credentials, multi-factor authentication, and role-based permissions. Passwords must be complex and changed regularly. Shared accounts are prohibited.",
        "Data classification identifies information sensitivity levels: public, internal, confidential, and restricted. Handling requirements vary by classification with stricter controls for sensitive data.",
        "Physical security measures protect facilities and equipment. Badge access, visitor logging, and clean desk policies prevent unauthorized physical access to sensitive areas.",
        "Network security employs firewalls, intrusion detection, and regular vulnerability assessments. Remote access requires VPN connections and endpoint protection software.",
        "Device security mandates encryption for laptops and mobile devices, automatic screen locking, and approved security software. Personal devices require explicit approval for work use.",
        "Incident response procedures ensure rapid detection and containment of security breaches. All suspicious activity must be reported immediately through established channels.",
        "Training requirements include annual security awareness sessions and specialized training for high-risk roles. Employees must demonstrate understanding of security policies.",
        "Third-party vendors must meet security standards before access is granted. Contracts include security requirements and regular audits ensure compliance.",
        "Continuous improvement involves regular security assessments, policy updates, and adoption of emerging security technologies. Employee feedback helps identify potential vulnerabilities.",
        "Remote work security requires secure Wi-Fi, updated software, and awareness of home network risks. Company-provided equipment offers additional protection layers.",
        "Data disposal follows secure practices including shredding, degaussing, and certified destruction services. Temporary files and caches are regularly cleared."
    ]

def generate_benefits_policy():
    """Generate benefits policy content."""
    return [
        "TechCorp Solutions offers comprehensive benefits designed to support employee health, financial security, and work-life balance. This policy outlines available benefits and enrollment procedures.",
        "Health insurance includes medical, dental, and vision coverage with multiple plan options. Premiums are shared between employer and employee, with the company covering approximately 80% of costs.",
        "Retirement planning features a 401(k) plan with employer matching up to 6% of salary. Employees can contribute pre-tax dollars with immediate vesting on all contributions.",
        "Life insurance provides coverage equal to annual salary, with options for additional voluntary coverage. Long-term disability insurance protects income in case of extended illness.",
        "Flexible spending accounts allow pre-tax contributions for medical and dependent care expenses. These accounts help reduce taxable income and cover eligible healthcare costs.",
        "Paid time off includes vacation days, sick leave, and holiday pay as outlined in the separate PTO policy. Accrual rates increase with tenure and performance.",
        "Professional development funding supports training, conferences, and education expenses up to $2,000 annually. Tuition reimbursement is available for job-related courses.",
        "Wellness programs encourage healthy lifestyles through gym memberships, health screenings, and wellness challenges. Employee assistance programs provide confidential counseling services.",
        "Parental benefits include paid maternity and paternity leave, with options for extended unpaid leave. Adoption assistance provides financial support for qualified adoptions.",
        "Employee discounts offer savings on various products and services through partnerships with local and national vendors. TechCorp branded merchandise is available at reduced prices.",
        "Commuter benefits help offset transportation costs with pre-tax contributions for public transit, parking, or vanpool services.",
        "Regular benefits fairs and one-on-one consultations ensure employees understand and maximize their benefits. Open enrollment occurs annually with opportunities for changes."
    ]

def main():
    """Generate all documents."""
    documents = [
        ("policy.pdf", "Employee Termination Policy", generate_termination_policy()),
        ("remote_work.pdf", "Remote Work Policy", generate_remote_work_policy()),
        ("holiday_policy.pdf", "Holiday and Time Off Policy", generate_holiday_policy()),
        ("code_of_conduct.pdf", "Employee Code of Conduct", generate_code_of_conduct()),
        ("security_guidelines.docx", "Information Security Guidelines", generate_security_guidelines()),
        ("benefits.pdf", "Employee Benefits Policy", generate_benefits_policy()),
    ]

    os.makedirs("data", exist_ok=True)

    for filename, title, paragraphs in documents:
        filepath = os.path.join("data", filename)

        if filename.endswith('.pdf'):
            create_pdf(filepath, title, paragraphs)
        elif filename.endswith('.docx'):
            create_docx(filepath, title, paragraphs)

        print(f"Generated: {filepath} - {len(paragraphs)} paragraphs")

if __name__ == "__main__":
    main()
